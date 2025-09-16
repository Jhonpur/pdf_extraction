
from docx import Document

loc2 = r"C:\Users\ange.kadjafomekon\OneDrive - AGM Solutions\Desktop\git_locale\pdf_extraction\Azienda 200 dip 03 - mese 01.2025.docx"
# Estrazione del testo paragrafo per paragrafo
doc = Document(loc2)
for para in doc.paragraphs:
    print(para.text)

# Estrazione delle tabelle (se presenti)
"""for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        print(row_data)"""


